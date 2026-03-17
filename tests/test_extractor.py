"""Tests for extractor - PDF 文本提取"""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from dockit.core.extractor import extract_text, MIN_TEXT_LENGTH


def test_extract_nonexistent():
    with pytest.raises(FileNotFoundError):
        extract_text("nonexistent_file_that_does_not_exist.pdf")


def test_extract_non_pdf(tmp_path):
    (tmp_path / "x.txt").write_text("hello")
    with pytest.raises(ValueError, match="不支持"):
        extract_text(tmp_path / "x.txt")


def test_extract_short_text_returns_invalid(tmp_path):
    """文本过短应返回 (text, False)"""
    pdf = tmp_path / "short.pdf"
    pdf.write_bytes(b"%PDF-1.4 dummy")
    with patch("dockit.core.extractor.pdfplumber") as mock_plumber:
        mock_pdf = MagicMock()
        mock_pdf.pages = [MagicMock(extract_text=MagicMock(return_value="Hi"))]
        mock_plumber.open.return_value.__enter__.return_value = mock_pdf
        text, valid = extract_text(pdf)
    assert valid is False
    assert len(text) < MIN_TEXT_LENGTH


def test_extract_enough_text_returns_valid(tmp_path):
    """文本足够长应返回 (text, True)"""
    content = "A" * (MIN_TEXT_LENGTH + 10)
    pdf = tmp_path / "ok.pdf"
    pdf.write_bytes(b"%PDF-1.4 dummy")
    with patch("dockit.core.extractor.pdfplumber") as mock_plumber:
        mock_pdf = MagicMock()
        mock_pdf.pages = [MagicMock(extract_text=MagicMock(return_value=content))]
        mock_plumber.open.return_value.__enter__.return_value = mock_pdf
        text, valid = extract_text(pdf)
    assert valid is True
    assert len(text) >= MIN_TEXT_LENGTH


def test_extract_docx_short_returns_invalid(tmp_path):
    """docx 文本过短返回 (text, False)"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("short")
    p = tmp_path / "short.docx"
    doc.save(p)
    text, valid = extract_text(p)
    assert valid is False


def test_extract_docx(tmp_path):
    """Word 文档提取"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("原告张三，被告李四。案由合同纠纷。" * 5)
    docx_path = tmp_path / "test.docx"
    doc.save(docx_path)
    text, valid = extract_text(docx_path)
    assert valid is True
    assert "原告" in text


def test_extract_pdf_exception_reraises(tmp_path):
    """PDF 提取失败时向上抛出异常"""
    pdf = tmp_path / "bad.pdf"
    pdf.write_bytes(b"%PDF-1.4 x")
    with patch("dockit.core.extractor.pdfplumber.open") as mock_open:
        mock_open.side_effect = RuntimeError("pdfplumber failed")
        with pytest.raises(RuntimeError, match="pdfplumber failed"):
            extract_text(pdf)


def test_extract_image_ocr_exception_returns_empty(tmp_path):
    """图片 OCR 异常时 _extract_from_image 返回空"""
    from dockit.core.extractor import _extract_from_image

    img = tmp_path / "x.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82")
    with patch("pytesseract.image_to_string", side_effect=RuntimeError("OCR failed")):
        result = _extract_from_image(img)
    assert result == ""


def test_extract_image_unsupported_ext_in_extract_text(tmp_path):
    """extract_text 对非支持扩展名抛出 ValueError"""
    (tmp_path / "x.gif").write_bytes(b"GIF89a")
    with pytest.raises(ValueError, match="不支持"):
        extract_text(tmp_path / "x.gif")
